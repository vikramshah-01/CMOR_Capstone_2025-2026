# Create Google Calendar import CSV and a Summary Checklist DOCX

import pandas as pd
from datetime import datetime

# -----------------------------
# 1. GOOGLE CALENDAR CSV FILE
# -----------------------------

events = [
    # HUMA 314
    ("HUMA 314 - Paper 1 Topic Paragraph Due", "02/15/2026", "11:00 PM", "11:59 PM"),
    ("HUMA 314 - Paper 1 Presentations", "03/10/2026", "09:00 AM", "10:00 AM"),
    ("HUMA 314 - Paper 1 Presentations", "03/12/2026", "09:00 AM", "10:00 AM"),
    ("HUMA 314 - Paper 1 Due", "03/15/2026", "11:00 PM", "11:59 PM"),
    ("HUMA 314 - Paper 2 Posted", "03/24/2026", "09:00 AM", "10:00 AM"),
    ("HUMA 314 - Paper 2 Presentations", "04/21/2026", "09:00 AM", "10:00 AM"),
    ("HUMA 314 - Paper 2 Presentations", "04/23/2026", "09:00 AM", "10:00 AM"),

    # RELI 101
    ("RELI 101 - Invent Your Own Religion Due", "03/24/2026", "11:00 PM", "11:59 PM"),
    ("RELI 101 - Final Exam (Part 1)", "04/21/2026", "04:00 PM", "05:15 PM"),
    ("RELI 101 - Final Exam (Part 2)", "04/23/2026", "04:00 PM", "05:15 PM"),

    # CMOR 544
    ("CMOR 544 - Final Project Presentation Option 1", "04/20/2026", "04:00 PM", "05:15 PM"),
    ("CMOR 544 - Final Project Presentation Option 2", "04/22/2026", "04:00 PM", "05:15 PM"),

    # CMOR 493
    ("CMOR 493 - Final Report & Presentation", "04/22/2026", "12:00 PM", "01:00 PM"),

    # Key Discussions
    ("CMOR 493 - Chapter 5 & 6 Discussion", "03/25/2026", "12:00 PM", "01:00 PM"),
    ("CMOR 493 - Chapter 7 Discussion", "04/08/2026", "12:00 PM", "01:00 PM"),
    ("CMOR 493 - Chapter 8 Discussion", "04/22/2026", "12:00 PM", "01:00 PM"),
]

calendar_df = pd.DataFrame(events, columns=["Subject", "Start Date", "Start Time", "End Time"])
calendar_df["End Date"] = calendar_df["Start Date"]
calendar_df["All Day Event"] = "False"
calendar_df["Description"] = ""

calendar_df = calendar_df[["Subject", "Start Date", "Start Time", "End Date", "End Time", "All Day Event", "Description"]]

calendar_path = "C:\\Users\\19015\Downloads\\Syllabi\\Spring_2026_Google_Calendar_Import.csv"
calendar_df.to_csv(calendar_path, index=False)

# -----------------------------
# 2. SUMMARY CHECKLIST DOCX
# -----------------------------

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
doc.add_heading("Spring 2026 Semester Summary Checklist", level=1)

sections = {
    "HUMA 314": [
        "☐ Paper 1 Topic Paragraph (Feb 15)",
        "☐ Paper 1 Presentations (Mar 10 & 12)",
        "☐ Paper 1 Final Submission (Mar 15)",
        "☐ Paper 2 Assigned (Mar 24)",
        "☐ Paper 2 Presentations (Apr 21 & 23)",
        "☐ Midterm (Late Feb)",
        "☐ Final Exam (Late March Release)"
    ],
    "RELI 101": [
        "☐ Invent Your Own Religion Report (Mar 24)",
        "☐ Read & Vote on Religions (Post-Spring Break)",
        "☐ 2 Vocabulary Pop Quizzes (TBD)",
        "☐ Final Exam (Apr 21 & 23)"
    ],
    "CMOR 544": [
        "☐ Homework (Ongoing - 5 Grace Days Available)",
        "☐ Individual Reading Report (Final Project)",
        "☐ Group Project Presentation (Apr 20 or 22)",
        "☐ Final Written Project Submission"
    ],
    "CMOR 493": [
        "☐ Chapter Discussions (Jan–Apr)",
        "☐ Final Report Submission",
        "☐ Final 10-min Presentation (Apr 22)"
    ],
    "CMOR 404": [
        "☐ Biweekly Problem Sets",
        "☐ Overleaf Research Paper (Biweekly Checks)",
        "☐ Dataset & Reproducibility Portfolio",
        "☐ Final Paper + Referee Revision Cycle"
    ]
}

for section, items in sections.items():
    doc.add_heading(section, level=2)
    for item in items:
        doc.add_paragraph(item)

checklist_path = "C:\\Users\\19015\\Downloads\\Syllabi\\Spring_2026_Semester_Checklist.docx"
doc.save(checklist_path)

calendar_path, checklist_path
